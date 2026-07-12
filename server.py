# -*- coding: utf-8 -*-
"""
מערכת ניהול תא הסברה — נפת הרטוב
Backend: Flask + SQLite. ריכוז שאלות, מסמכים, מסרים, פעילות וסיכומי משמרת.
"""
import json
import os
import re
import secrets
import sqlite3
import string
import uuid
import hmac
import zipfile
from datetime import datetime, timedelta
from functools import wraps
from zoneinfo import ZoneInfo

from flask import Flask, request, session, jsonify, send_from_directory, send_file, Response
from werkzeug.security import generate_password_hash, check_password_hash

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# ב-Render: DATA_DIR מצביע לדיסק מתמיד; מקומית — תיקיית data בפרויקט
DATA_ROOT = os.environ.get('DATA_DIR') or os.path.join(BASE_DIR, 'data')
DB_PATH = os.path.join(DATA_ROOT, 'hasbara.db')
UPLOAD_DIR = os.path.join(DATA_ROOT, 'uploads')
os.makedirs(UPLOAD_DIR, exist_ok=True)

IS_PROD = bool(os.environ.get('DATA_DIR'))

app = Flask(__name__, static_folder='static')
app.secret_key = os.environ.get('SECRET_KEY', 'hasbara-cell-dev-key')
app.config.update(
    SESSION_COOKIE_HTTPONLY=True,
    SESSION_COOKIE_SAMESITE='Lax',
    SESSION_COOKIE_SECURE=IS_PROD,
    PERMANENT_SESSION_LIFETIME=timedelta(hours=12),
    MAX_CONTENT_LENGTH=20 * 1024 * 1024,   # 20MB — העלאת מסמכים
)

ANTHROPIC_API_KEY = os.environ.get('ANTHROPIC_API_KEY', '')
DEFAULT_MODEL = os.environ.get('ANTHROPIC_MODEL', 'claude-sonnet-5')
WEBHOOK_TOKEN = os.environ.get('WHATSAPP_WEBHOOK_TOKEN', '')

# ---------------------------------------------------------------- קבועים
STATUSES = ['חדש', 'בטיפול', 'ממתין למידע', 'ממתין לאישור', 'נענה', 'נסגר', 'דורש המשך טיפול']
OPEN_STATUSES = ['חדש', 'בטיפול', 'ממתין למידע', 'ממתין לאישור', 'דורש המשך טיפול']
URGENCIES = ['רגיל', 'דחוף', 'מיידי']
SOURCES = ['ידני', 'וואטסאפ', 'מייל', 'מסמך', 'אחר']
DOC_TYPES = ['הנחיית פיקוד העורף', 'סיכום מצב', 'הודעה רשמית', 'נוהל', 'מסמך רשות',
             'מסמך פנימי', 'מסמך מחוז / נפה', 'עדכון מדיניות', 'אחר']
MSG_STATUSES = ['טיוטה', 'ממתין לאישור', 'פעיל', 'לא פעיל', 'הוחלף']
AUDIENCES = ['רשויות', 'יקל"רים', 'דוברים', 'מוקד', 'ציבור', 'בעלי תפקידים', 'הנהלה / חמ"ל']
ACTIVITY_TYPES = ['הודעה שהופצה', 'מענה לשאלה', 'עדכון לרשויות', 'תדרוך', 'שיחה עם גורם שטח',
                  'חומר לרשתות', 'מסמך מסרים', 'סיכום מצב', 'הפצת הנחיה', 'עדכון למוקד',
                  'עדכון לדוברים', 'אחר']
ACTIVITY_STATUSES = ['בוצע', 'בטיפול', 'ממתין לאישור', 'הופץ', 'דורש המשך טיפול']
OUT_STATUSES = ['טיוטה', 'מאושר', 'הופץ']
DISTRICT_NAME = os.environ.get('DISTRICT_NAME', 'נפת הרטוב')
CANNED_CATEGORIES = ['ירי רקטות וטילים', 'רעידת אדמה', 'חומרים מסוכנים', 'אירוע ביטחוני',
                     'מזג אוויר קיצון', 'הרגעה ועדכון כללי', 'אחר']
MATERIAL_CATEGORIES = ['פלייר', 'אינפוגרפיקה', 'תמונה', 'סרטון', 'מצגת', 'מסמך', 'אחר']
MATERIAL_EXTS = ('.png', '.jpg', '.jpeg', '.gif', '.webp', '.pdf', '.mp4', '.pptx', '.docx', '.xlsx')
ROLES = ['admin', 'lead', 'user', 'viewer']
ROLE_NAMES = {'admin': 'מנהל מערכת', 'lead': 'אחראי תא הסברה', 'user': 'משתמש', 'viewer': 'צפייה בלבד'}

SEED_TOPICS = ['מדיניות התגוננות', 'לימודים', 'התקהלות', 'מקלטים', 'שירותים חיוניים',
               'מידע לציבור', 'הנחיות פיקוד העורף', 'פעילות רשותית', 'פערי מידע', 'אחר']

# בהרצה ראשונה נוצר משתמש admin בלבד; הסיסמה מ-ADMIN_PASSWORD (ב-Render נוצר
# אוטומטית ומוצג ב-Environment) או אקראית שמודפסת לקונסול. שאר המשתמשים
# נוצרים ממסך ההגדרות.


# שעון ישראל — שרתי Render רצים ב-UTC; כל חותמות הזמן במערכת בשעון ישראל
IL_TZ = ZoneInfo('Asia/Jerusalem')


def now_dt():
    return datetime.now(IL_TZ).replace(tzinfo=None)


def now():
    return now_dt().isoformat(timespec='seconds')


def today():
    return now_dt().date().isoformat()


# ---------------------------------------------------------------- DB
def get_db():
    conn = sqlite3.connect(DB_PATH, timeout=10)
    conn.row_factory = sqlite3.Row
    conn.execute('PRAGMA foreign_keys = ON')
    conn.execute('PRAGMA journal_mode = WAL')
    conn.execute('PRAGMA busy_timeout = 5000')
    return conn


def init_db():
    conn = get_db()
    c = conn.cursor()
    c.executescript('''
        CREATE TABLE IF NOT EXISTS users (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL,
            username TEXT UNIQUE NOT NULL,
            password_hash TEXT NOT NULL,
            role TEXT NOT NULL DEFAULT 'user',
            active INTEGER NOT NULL DEFAULT 1,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS questions (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            opened_at TEXT NOT NULL,
            asker_name TEXT,
            asker_phone TEXT,
            source TEXT NOT NULL DEFAULT 'ידני',
            authority TEXT,
            content TEXT NOT NULL,
            topic TEXT,
            urgency TEXT NOT NULL DEFAULT 'רגיל',
            assignee_id INTEGER,
            status TEXT NOT NULL DEFAULT 'חדש',
            proposed_answer TEXT,
            approved_answer TEXT,
            needs_approval INTEGER NOT NULL DEFAULT 0,
            approved_by_id INTEGER,
            answered_at TEXT,
            closed_at TEXT,
            internal_notes TEXT,
            raw_source_text TEXT,
            created_by_id INTEGER,
            updated_at TEXT
        );
        CREATE TABLE IF NOT EXISTS question_history (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_id INTEGER NOT NULL,
            action TEXT NOT NULL,
            detail TEXT,
            user_id INTEGER,
            created_at TEXT NOT NULL,
            FOREIGN KEY (question_id) REFERENCES questions(id)
        );
        CREATE TABLE IF NOT EXISTS attachments (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_id INTEGER NOT NULL,
            orig_name TEXT NOT NULL,
            stored_name TEXT NOT NULL,
            mime TEXT,
            size INTEGER,
            uploaded_by_id INTEGER,
            created_at TEXT NOT NULL,
            FOREIGN KEY (question_id) REFERENCES questions(id)
        );
        CREATE TABLE IF NOT EXISTS question_links (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            question_id INTEGER NOT NULL,
            kind TEXT NOT NULL,               -- document / message / activity
            ref_id INTEGER NOT NULL,
            created_at TEXT NOT NULL,
            FOREIGN KEY (question_id) REFERENCES questions(id)
        );
        CREATE TABLE IF NOT EXISTS documents (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            doc_type TEXT,
            source TEXT,
            orig_name TEXT,
            stored_name TEXT,
            mime TEXT,
            size INTEGER,
            extracted_text TEXT,
            ai_summary TEXT,
            ai_key_points TEXT,
            ai_messages TEXT,
            ai_qa TEXT,
            ai_gaps TEXT,
            ai_draft_message TEXT,
            insights_generated_at TEXT,
            insights_model TEXT,
            uploaded_by_id INTEGER,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            body TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'טיוטה',
            audience TEXT,
            valid_from TEXT,
            valid_until TEXT,
            source_document_id INTEGER,
            replaced_by_id INTEGER,
            created_by_id INTEGER,
            approved_by_id INTEGER,
            notes TEXT,
            created_at TEXT NOT NULL,
            updated_at TEXT
        );
        CREATE TABLE IF NOT EXISTS activities (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            activity_type TEXT NOT NULL,
            description TEXT NOT NULL,
            topic TEXT,
            authority TEXT,
            audience TEXT,
            status TEXT NOT NULL DEFAULT 'בוצע',
            performed_at TEXT NOT NULL,
            performed_by_id INTEGER,
            question_id INTEGER,
            document_id INTEGER,
            message_id INTEGER,
            notes TEXT,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS outgoing_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            source_type TEXT NOT NULL,        -- question / document / message / manual
            source_id INTEGER,
            audience TEXT,
            body TEXT NOT NULL,
            status TEXT NOT NULL DEFAULT 'טיוטה',
            needs_approval INTEGER NOT NULL DEFAULT 0,
            created_by_id INTEGER,
            approved_by_id INTEGER,
            distributed_at TEXT,
            activity_id INTEGER,
            ai_generated INTEGER NOT NULL DEFAULT 0,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS shift_summaries (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            period_start TEXT NOT NULL,
            period_end TEXT NOT NULL,
            body TEXT NOT NULL,
            stats_json TEXT,
            ai_generated INTEGER NOT NULL DEFAULT 0,
            created_by_id INTEGER,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS topics (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL,
            active INTEGER NOT NULL DEFAULT 1,
            sort_order INTEGER NOT NULL DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS authorities (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT UNIQUE NOT NULL,
            active INTEGER NOT NULL DEFAULT 1,
            sort_order INTEGER NOT NULL DEFAULT 0
        );
        CREATE TABLE IF NOT EXISTS canned_messages (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            category TEXT,
            body TEXT NOT NULL,
            audience TEXT,
            notes TEXT,
            active INTEGER NOT NULL DEFAULT 1,
            created_by_id INTEGER,
            created_at TEXT NOT NULL,
            updated_at TEXT
        );
        CREATE TABLE IF NOT EXISTS materials (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            title TEXT NOT NULL,
            category TEXT,
            description TEXT,
            orig_name TEXT NOT NULL,
            stored_name TEXT NOT NULL,
            mime TEXT,
            size INTEGER,
            uploaded_by_id INTEGER,
            created_at TEXT NOT NULL
        );
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT
        );
        CREATE TABLE IF NOT EXISTS audit_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            user_id INTEGER,
            action TEXT NOT NULL,
            entity TEXT,
            entity_id INTEGER,
            detail TEXT,
            created_at TEXT NOT NULL
        );
        CREATE INDEX IF NOT EXISTS idx_questions_status ON questions(status);
        CREATE INDEX IF NOT EXISTS idx_questions_opened ON questions(opened_at);
        CREATE INDEX IF NOT EXISTS idx_activities_performed ON activities(performed_at);
        CREATE INDEX IF NOT EXISTS idx_qhistory_qid ON question_history(question_id);
    ''')
    conn.commit()

    # מיגרציות עתידיות לעמודות חדשות (DB קיים)
    def add_col(table, col, ddl):
        cols = [r['name'] for r in c.execute(f'PRAGMA table_info({table})').fetchall()]
        if col not in cols:
            c.execute(f'ALTER TABLE {table} ADD COLUMN {ddl}')
    # (אין מיגרציות כרגע — התבנית מוכנה לעתיד)

    if c.execute('SELECT COUNT(*) FROM users').fetchone()[0] == 0:
        admin_pw = os.environ.get('ADMIN_PASSWORD') or ''.join(
            secrets.choice(string.ascii_letters + string.digits) for _ in range(12))
        c.execute('INSERT INTO users (name, username, password_hash, role, active, created_at) '
                  'VALUES (?,?,?,?,1,?)',
                  ('מנהל מערכת', 'admin', generate_password_hash(admin_pw), 'admin', now()))
        if not os.environ.get('ADMIN_PASSWORD'):
            print(f'*** משתמש ראשוני: admin | סיסמה: {admin_pw} — החלף מיד לאחר הכניסה ***')
    if c.execute('SELECT COUNT(*) FROM topics').fetchone()[0] == 0:
        for i, t in enumerate(SEED_TOPICS):
            c.execute('INSERT INTO topics (name, active, sort_order) VALUES (?,1,?)', (t, i))
    # תבניות הודעה ראשוניות לבנק — ניתנות לעריכה/מחיקה מהמסך
    if c.execute('SELECT COUNT(*) FROM canned_messages').fetchone()[0] == 0:
        seeds = [
            ('הודעה ראשונית — ירי רקטות', 'ירי רקטות וטילים', 'ציבור',
             'בעקבות ירי לעבר [אזור]: יש לפעול לפי הנחיות פיקוד העורף. בהישמע אזעקה — היכנסו '
             'למרחב המוגן ושהו בו [10] דקות. עקבו אחר עדכונים בערוצים הרשמיים בלבד.'),
            ('הודעה ראשונית — רעידת אדמה', 'רעידת אדמה', 'ציבור',
             'בעקבות רעידת האדמה שהורגשה ב-[שעה]: אם אתם בתוך מבנה — צאו לשטח פתוח. '
             'התרחקו ממבנים, עצים ועמודי חשמל. אין להשתמש במעליות. המוקד העירוני: [מספר].'),
            ('הודעה ראשונית — אירוע חומרים מסוכנים', 'חומרים מסוכנים', 'ציבור',
             'בעקבות אירוע חומרים מסוכנים ב-[מיקום]: תושבי [אזור] מתבקשים להיכנס למבנה, '
             'לסגור חלונות ולכבות מזגנים עד להודעה חדשה. אין להתקרב לאזור האירוע.'),
            ('עדכון הרגעה כללי', 'הרגעה ועדכון כללי', 'ציבור',
             'עדכון לתושבי [אזור]: האירוע ב-[מיקום] בטיפול הכוחות. אין הנחיות מיוחדות לציבור '
             'בשלב זה. נעדכן בכל שינוי בערוצים הרשמיים.'),
        ]
        for title, cat, aud, body in seeds:
            c.execute('INSERT INTO canned_messages (title, category, body, audience, active, created_at) '
                      'VALUES (?,?,?,?,1,?)', (title, cat, body, aud, now()))
    conn.commit()
    conn.close()


def get_setting(conn, key, default=None):
    row = conn.execute('SELECT value FROM settings WHERE key=?', (key,)).fetchone()
    return row['value'] if row else default


def set_setting(conn, key, value):
    conn.execute('INSERT INTO settings (key, value) VALUES (?,?) '
                 'ON CONFLICT(key) DO UPDATE SET value=excluded.value', (key, value))


def log_action(conn, action, entity=None, entity_id=None, detail=None):
    conn.execute('INSERT INTO audit_log (user_id, action, entity, entity_id, detail, created_at) '
                 'VALUES (?,?,?,?,?,?)',
                 (session.get('uid'), action, entity, entity_id, detail, now()))


def log_history(conn, qid, action, detail=None, user_id=None):
    conn.execute('INSERT INTO question_history (question_id, action, detail, user_id, created_at) '
                 'VALUES (?,?,?,?,?)',
                 (qid, action, detail, user_id if user_id is not None else session.get('uid'), now()))


def user_name(conn, uid):
    if not uid:
        return None
    row = conn.execute('SELECT name FROM users WHERE id=?', (uid,)).fetchone()
    return row['name'] if row else None


# ---------------------------------------------------------------- auth
def roles_required(*roles):
    def deco(f):
        @wraps(f)
        def wrapper(*a, **k):
            if session.get('role') not in roles:
                return jsonify({'error': 'אין הרשאה לפעולה זו'}), 403
            return f(*a, **k)
        return wrapper
    return deco


PUBLIC_PATHS = ('/api/login', '/api/webhook/whatsapp', '/api/webhook/greenapi')


@app.before_request
def api_gate():
    p = request.path
    if not p.startswith('/api/'):
        return None
    if p in PUBLIC_PATHS:
        return None
    if 'uid' not in session:
        return jsonify({'error': 'לא מחובר'}), 401
    # צפייה בלבד — קריאה בלבד (מלבד יציאה ושינוי סיסמה עצמי)
    if request.method != 'GET' and session.get('role') == 'viewer' \
            and p not in ('/api/logout', '/api/change-password'):
        return jsonify({'error': 'הרשאת צפייה בלבד'}), 403
    return None


def webhook_token_ok():
    """אימות טוקן ה-webhook. תו '+' בטוקן מפוענח כרווח ב-query string — מקבלים את שתי הצורות."""
    supplied = request.args.get('token', '')
    if not WEBHOOK_TOKEN or not supplied:
        return False
    return (hmac.compare_digest(supplied, WEBHOOK_TOKEN)
            or hmac.compare_digest(supplied.replace(' ', '+'), WEBHOOK_TOKEN))


# הגבלת קצב פשוטה נגד ניחוש סיסמאות (בזיכרון; מספיק למופע יחיד)
_login_fails = {}
LOGIN_MAX = 10
LOGIN_WINDOW = 300


def _too_many_attempts(ip):
    ts = datetime.now().timestamp()
    fails = [t for t in _login_fails.get(ip, []) if ts - t < LOGIN_WINDOW]
    _login_fails[ip] = fails
    return len(fails) >= LOGIN_MAX


@app.route('/api/login', methods=['POST'])
def login():
    ip = request.headers.get('X-Forwarded-For', request.remote_addr or '').split(',')[0].strip()
    if _too_many_attempts(ip):
        return jsonify({'error': 'יותר מדי ניסיונות. נסה שוב בעוד כמה דקות.'}), 429
    data = request.get_json(force=True)
    conn = get_db()
    row = conn.execute('SELECT * FROM users WHERE username=? AND active=1',
                       (data.get('username', '').strip(),)).fetchone()
    conn.close()
    if row and check_password_hash(row['password_hash'], data.get('password', '')):
        session.clear()
        session.permanent = True
        session['uid'] = row['id']
        session['role'] = row['role']
        session['name'] = row['name']
        return jsonify({'id': row['id'], 'name': row['name'], 'role': row['role']})
    _login_fails.setdefault(ip, []).append(datetime.now().timestamp())
    return jsonify({'error': 'שם משתמש או סיסמה שגויים'}), 401


@app.route('/api/logout', methods=['POST'])
def logout():
    session.clear()
    return jsonify({'ok': True})


@app.route('/api/me')
def me():
    return jsonify({'id': session['uid'], 'name': session['name'], 'role': session['role']})


@app.route('/api/change-password', methods=['POST'])
def change_password():
    d = request.get_json(force=True)
    new = d.get('new') or ''
    if len(new) < 8:
        return jsonify({'error': 'הסיסמה החדשה חייבת לפחות 8 תווים'}), 400
    conn = get_db()
    row = conn.execute('SELECT * FROM users WHERE id=?', (session['uid'],)).fetchone()
    if not row or not check_password_hash(row['password_hash'], d.get('current', '')):
        conn.close()
        return jsonify({'error': 'הסיסמה הנוכחית שגויה'}), 403
    conn.execute('UPDATE users SET password_hash=? WHERE id=?',
                 (generate_password_hash(new), session['uid']))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ---------------------------------------------------------------- meta
@app.route('/api/meta')
def meta():
    conn = get_db()
    topics = [r['name'] for r in conn.execute(
        'SELECT name FROM topics WHERE active=1 ORDER BY sort_order, name')]
    authorities = [r['name'] for r in conn.execute(
        'SELECT name FROM authorities WHERE active=1 ORDER BY sort_order, name')]
    users = [dict(r) for r in conn.execute(
        'SELECT id, name, role FROM users WHERE active=1 ORDER BY name')]
    conn.close()
    return jsonify({
        'topics': topics, 'authorities': authorities, 'users': users,
        'statuses': STATUSES, 'urgencies': URGENCIES, 'sources': SOURCES,
        'doc_types': DOC_TYPES, 'msg_statuses': MSG_STATUSES, 'audiences': AUDIENCES,
        'activity_types': ACTIVITY_TYPES, 'activity_statuses': ACTIVITY_STATUSES,
        'out_statuses': OUT_STATUSES,
        'canned_categories': CANNED_CATEGORIES, 'material_categories': MATERIAL_CATEGORIES,
        'ai_enabled': ai_enabled(),
        'wa_send_enabled': wa_send_enabled(),
    })


# ---------------------------------------------------------------- Claude AI
def get_api_key():
    """מפתח API: קודם מהגדרות המערכת (DB), אחרת ממשתנה סביבה."""
    conn = get_db()
    key = get_setting(conn, 'anthropic_api_key')
    conn.close()
    return key or ANTHROPIC_API_KEY


def ai_enabled():
    return bool(get_api_key())


def get_model():
    conn = get_db()
    m = get_setting(conn, 'anthropic_model')
    conn.close()
    return m or DEFAULT_MODEL


def call_claude(system, user_text, max_tokens=4000, schema=None):
    """קריאה ל-Claude. מחזיר טקסט. זורק חריגות anthropic — לטפל בקורא."""
    import anthropic
    client = anthropic.Anthropic(api_key=get_api_key())
    kwargs = {}
    if schema:
        kwargs['output_config'] = {'format': {'type': 'json_schema', 'schema': schema}}
    msg = client.messages.create(
        model=get_model(),
        max_tokens=max_tokens,
        system=system,
        messages=[{'role': 'user', 'content': user_text}],
        **kwargs)
    return next((b.text for b in msg.content if b.type == 'text'), '')


def ai_error(e):
    """מיפוי חריגות anthropic להודעה בעברית."""
    import anthropic
    if isinstance(e, anthropic.AuthenticationError):
        msg = 'מפתח ה-API אינו תקין'
    elif isinstance(e, anthropic.RateLimitError):
        msg = 'עומס על שירות ה-AI — נסה שוב בעוד רגע'
    elif isinstance(e, anthropic.APIConnectionError):
        msg = 'אין חיבור לשירות ה-AI'
    elif isinstance(e, anthropic.APIStatusError):
        msg = f'שגיאת שירות AI ({e.status_code})'
    else:
        msg = 'שגיאה בקריאה לשירות ה-AI'
    return jsonify({'error': msg}), 502


INSIGHTS_SCHEMA = {
    'type': 'object',
    'properties': {
        'summary': {'type': 'string', 'description': 'תקציר קצר של המסמך בעברית'},
        'key_points': {'type': 'string', 'description': 'עיקרי הדברים, שורה לכל נקודה'},
        'messages': {'type': 'string', 'description': 'מסרים מרכזיים לציבור, שורה לכל מסר'},
        'qa': {'type': 'string', 'description': 'שאלות ותשובות צפויות מהשטח, בפורמט ש: ... ת: ...'},
        'gaps': {'type': 'string', 'description': 'פערי מידע ונקודות לא ברורות שדורשות הבהרה'},
        'draft_message': {'type': 'string', 'description': 'נוסח הודעה קצרה מוצעת להפצה'},
    },
    'required': ['summary', 'key_points', 'messages', 'qa', 'gaps', 'draft_message'],
    'additionalProperties': False,
}

INSIGHTS_SYSTEM = (
    f'אתה קצין הסברה ב{DISTRICT_NAME} של פיקוד העורף. תפקידך לנתח מסמכים ולהפוך אותם לחומר עבודה '
    'הסברתי מעשי. נתח את המסמך שתקבל והחזר JSON בעברית עם השדות: summary (תקציר קצר), '
    'key_points (עיקרי הדברים — שורה לכל נקודה, פתח כל שורה במקף), messages (מסרים מרכזיים '
    'לציבור — שורה לכל מסר, פתח כל שורה במקף), qa (שאלות ותשובות צפויות מהשטח בפורמט '
    '"ש: ... ת: ..."), gaps (פערי מידע שדורשים הבהרה), draft_message (נוסח הודעה קצרה '
    'וברורה להפצה לציבור). כתוב בעברית פשוטה וברורה, מתאימה לשעת חירום. '
    'אם הטקסט חלקי או משובש (חילוץ PDF), עשה כמיטב יכולתך והצנע זאת בפערים.'
)


# ---------------------------------------------------------------- שאלות
def question_row(conn, r):
    d = dict(r)
    d['assignee_name'] = user_name(conn, r['assignee_id'])
    d['approved_by_name'] = user_name(conn, r['approved_by_id'])
    d['created_by_name'] = user_name(conn, r['created_by_id'])
    return d


@app.route('/api/questions')
def list_questions():
    conn = get_db()
    where, params = [], []
    for field in ('status', 'topic', 'urgency', 'source', 'authority'):
        v = request.args.get(field)
        if v:
            where.append(f'{field}=?')
            params.append(v)
    if request.args.get('assignee'):
        where.append('assignee_id=?')
        params.append(request.args['assignee'])
    if request.args.get('open') == '1':
        where.append('status IN (%s)' % ','.join('?' * len(OPEN_STATUSES)))
        params.extend(OPEN_STATUSES)
    if request.args.get('from'):
        where.append('opened_at >= ?')
        params.append(request.args['from'])
    if request.args.get('to'):
        where.append('opened_at <= ?')
        params.append(request.args['to'] + 'T23:59:59')
    q = request.args.get('q', '').strip()
    if q:
        for term in q.split():
            where.append('(content LIKE ? OR proposed_answer LIKE ? OR approved_answer LIKE ? '
                         'OR asker_name LIKE ? OR internal_notes LIKE ?)')
            params.extend([f'%{term}%'] * 5)
    sql = 'SELECT * FROM questions'
    if where:
        sql += ' WHERE ' + ' AND '.join(where)
    sql += ' ORDER BY opened_at DESC LIMIT 500'
    rows = [question_row(conn, r) for r in conn.execute(sql, params)]
    conn.close()
    return jsonify(rows)


@app.route('/api/questions', methods=['POST'])
def create_question():
    d = request.get_json(force=True)
    content = (d.get('content') or '').strip()
    if not content:
        return jsonify({'error': 'תוכן השאלה חובה'}), 400
    conn = get_db()
    cur = conn.execute(
        'INSERT INTO questions (opened_at, asker_name, asker_phone, source, authority, content, '
        'topic, urgency, assignee_id, status, needs_approval, internal_notes, raw_source_text, '
        'created_by_id, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,?,?)',
        (d.get('opened_at') or now(), d.get('asker_name'), d.get('asker_phone'),
         d.get('source') or 'ידני', d.get('authority'), content, d.get('topic'),
         d.get('urgency') or 'רגיל', d.get('assignee_id') or None, d.get('status') or 'חדש',
         1 if d.get('needs_approval') else 0, d.get('internal_notes'),
         d.get('raw_source_text'), session['uid'], now()))
    qid = cur.lastrowid
    log_history(conn, qid, 'created', 'שאלה נפתחה')
    log_action(conn, 'create_question', 'question', qid)
    conn.commit()
    conn.close()
    return jsonify({'id': qid})


@app.route('/api/questions/<int:qid>')
def get_question(qid):
    conn = get_db()
    r = conn.execute('SELECT * FROM questions WHERE id=?', (qid,)).fetchone()
    if not r:
        conn.close()
        return jsonify({'error': 'שאלה לא נמצאה'}), 404
    d = question_row(conn, r)
    d['history'] = [dict(h) | {'user_name': user_name(conn, h['user_id'])}
                    for h in conn.execute(
                        'SELECT * FROM question_history WHERE question_id=? ORDER BY created_at DESC, id DESC',
                        (qid,))]
    d['attachments'] = [dict(a) for a in conn.execute(
        'SELECT id, orig_name, size, created_at FROM attachments WHERE question_id=?', (qid,))]
    links = []
    for l in conn.execute('SELECT * FROM question_links WHERE question_id=?', (qid,)):
        item = dict(l)
        if l['kind'] == 'document':
            ref = conn.execute('SELECT title FROM documents WHERE id=?', (l['ref_id'],)).fetchone()
            item['label'] = ref['title'] if ref else '?'
        elif l['kind'] == 'message':
            ref = conn.execute('SELECT title FROM messages WHERE id=?', (l['ref_id'],)).fetchone()
            item['label'] = ref['title'] if ref else '?'
        else:
            ref = conn.execute('SELECT description FROM activities WHERE id=?', (l['ref_id'],)).fetchone()
            item['label'] = (ref['description'][:60] if ref else '?')
        links.append(item)
    d['links'] = links
    conn.close()
    return jsonify(d)


EDITABLE_Q_FIELDS = ['asker_name', 'asker_phone', 'source', 'authority', 'content', 'topic',
                     'urgency', 'assignee_id', 'proposed_answer', 'needs_approval',
                     'internal_notes']
Q_FIELD_LABELS = {'asker_name': 'שם הפונה', 'asker_phone': 'טלפון', 'source': 'מקור',
                  'authority': 'רשות', 'content': 'תוכן', 'topic': 'נושא', 'urgency': 'דחיפות',
                  'assignee_id': 'גורם מטפל', 'proposed_answer': 'מענה מוצע',
                  'needs_approval': 'נדרש אישור', 'internal_notes': 'הערות פנימיות'}


@app.route('/api/questions/<int:qid>', methods=['PUT'])
def update_question(qid):
    d = request.get_json(force=True)
    conn = get_db()
    old = conn.execute('SELECT * FROM questions WHERE id=?', (qid,)).fetchone()
    if not old:
        conn.close()
        return jsonify({'error': 'שאלה לא נמצאה'}), 404
    changes, sets, params = [], [], []
    for f in EDITABLE_Q_FIELDS:
        if f in d:
            newv = d[f]
            if f == 'needs_approval':
                newv = 1 if newv else 0
            if newv != old[f]:
                sets.append(f'{f}=?')
                params.append(newv)
                if f == 'assignee_id':
                    changes.append(f'{Q_FIELD_LABELS[f]}: {user_name(conn, old[f]) or "-"} → {user_name(conn, newv) or "-"}')
                elif f in ('proposed_answer', 'internal_notes', 'content'):
                    changes.append(f'{Q_FIELD_LABELS[f]} עודכן')
                else:
                    changes.append(f'{Q_FIELD_LABELS[f]}: {old[f] or "-"} → {newv or "-"}')
    if sets:
        sets.append('updated_at=?')
        params.append(now())
        params.append(qid)
        conn.execute(f'UPDATE questions SET {", ".join(sets)} WHERE id=?', params)
        log_history(conn, qid, 'updated', '; '.join(changes))
        conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/questions/<int:qid>/status', methods=['POST'])
def set_question_status(qid):
    d = request.get_json(force=True)
    status = d.get('status')
    if status not in STATUSES:
        return jsonify({'error': 'סטטוס לא חוקי'}), 400
    conn = get_db()
    old = conn.execute('SELECT status FROM questions WHERE id=?', (qid,)).fetchone()
    if not old:
        conn.close()
        return jsonify({'error': 'שאלה לא נמצאה'}), 404
    closed = now() if status == 'נסגר' else None
    conn.execute('UPDATE questions SET status=?, closed_at=COALESCE(?, closed_at), updated_at=? WHERE id=?',
                 (status, closed, now(), qid))
    log_history(conn, qid, 'status_changed', f'{old["status"]} → {status}')
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/questions/<int:qid>/approve', methods=['POST'])
@roles_required('admin', 'lead')
def approve_question(qid):
    d = request.get_json(force=True)
    conn = get_db()
    q = conn.execute('SELECT * FROM questions WHERE id=?', (qid,)).fetchone()
    if not q:
        conn.close()
        return jsonify({'error': 'שאלה לא נמצאה'}), 404
    answer = (d.get('answer') or q['proposed_answer'] or '').strip()
    if not answer:
        conn.close()
        return jsonify({'error': 'אין מענה לאישור'}), 400
    conn.execute('UPDATE questions SET approved_answer=?, approved_by_id=?, status=?, '
                 'answered_at=?, updated_at=? WHERE id=?',
                 (answer, session['uid'], 'נענה', now(), now(), qid))
    log_history(conn, qid, 'approved', f'המענה אושר ע"י {session["name"]}')
    log_action(conn, 'approve_answer', 'question', qid)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/questions/<int:qid>/attachments', methods=['POST'])
def upload_attachment(qid):
    f = request.files.get('file')
    if not f or not f.filename:
        return jsonify({'error': 'לא נבחר קובץ'}), 400
    ext = os.path.splitext(f.filename)[1].lower()[:10]
    stored = uuid.uuid4().hex + ext
    f.save(os.path.join(UPLOAD_DIR, stored))
    size = os.path.getsize(os.path.join(UPLOAD_DIR, stored))
    conn = get_db()
    conn.execute('INSERT INTO attachments (question_id, orig_name, stored_name, mime, size, '
                 'uploaded_by_id, created_at) VALUES (?,?,?,?,?,?,?)',
                 (qid, f.filename, stored, f.mimetype, size, session['uid'], now()))
    log_history(conn, qid, 'attachment', f'צורף קובץ: {f.filename}')
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/attachments/<int:aid>/download')
def download_attachment(aid):
    conn = get_db()
    a = conn.execute('SELECT * FROM attachments WHERE id=?', (aid,)).fetchone()
    conn.close()
    if not a:
        return jsonify({'error': 'קובץ לא נמצא'}), 404
    return send_file(os.path.join(UPLOAD_DIR, a['stored_name']),
                     download_name=a['orig_name'], as_attachment=True)


@app.route('/api/questions/<int:qid>/links', methods=['POST'])
def add_link(qid):
    d = request.get_json(force=True)
    kind, ref_id = d.get('kind'), d.get('ref_id')
    if kind not in ('document', 'message', 'activity') or not ref_id:
        return jsonify({'error': 'קישור לא חוקי'}), 400
    conn = get_db()
    conn.execute('INSERT INTO question_links (question_id, kind, ref_id, created_at) VALUES (?,?,?,?)',
                 (qid, kind, ref_id, now()))
    log_history(conn, qid, 'linked', f'קושר {kind} #{ref_id}')
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/questions/<int:qid>/links/<int:lid>', methods=['DELETE'])
def delete_link(qid, lid):
    conn = get_db()
    conn.execute('DELETE FROM question_links WHERE id=? AND question_id=?', (lid, qid))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ---------------------------------------------------------------- וואטסאפ
WA_LINE_RE = re.compile(
    r'^\[?(\d{1,2}[./]\d{1,2}[./]\d{2,4}),?\s+(\d{1,2}:\d{2})(?::\d{2})?\]?\s*(?:-\s*)?([^:]{1,50}):\s*(.*)$')
PHONE_RE = re.compile(r'(?:\+?972[-\s]?|0)5\d[-\s]?\d{3}[-\s]?\d{4}')


def normalize_phone(p):
    if not p:
        return p
    p = re.sub(r'[\s\-]', '', p)
    if p.startswith('+972'):
        p = '0' + p[4:]
    elif p.startswith('972'):
        p = '0' + p[3:]
    return p


def parse_wa_date(dstr, tstr):
    try:
        parts = re.split(r'[./]', dstr)
        day, month, year = int(parts[0]), int(parts[1]), int(parts[2])
        if year < 100:
            year += 2000
        h, m = tstr.split(':')
        return datetime(year, month, day, int(h), int(m)).isoformat(timespec='seconds')
    except (ValueError, IndexError):
        return None


@app.route('/api/questions/paste-whatsapp', methods=['POST'])
def paste_whatsapp():
    """מקבל טקסט מודבק מוואטסאפ ומחזיר preview לאישור לפני יצירה."""
    text = (request.get_json(force=True).get('text') or '').strip()
    if not text:
        return jsonify({'error': 'לא הודבק טקסט'}), 400
    sender, opened_at, bodies = None, None, []
    for line in text.splitlines():
        m = WA_LINE_RE.match(line.strip())
        if m:
            if sender is None:
                sender = m.group(3).strip()
                opened_at = parse_wa_date(m.group(1), m.group(2))
            if m.group(3).strip() == sender:
                bodies.append(m.group(4))
        elif bodies:
            bodies.append(line)
    content = '\n'.join(bodies).strip() if bodies else text
    phone_m = PHONE_RE.search(text)
    phone = normalize_phone(phone_m.group(0)) if phone_m else None
    # אם ה"שם" שזוהה הוא בעצם מספר טלפון
    if sender and PHONE_RE.fullmatch(sender.replace(' ', '')):
        phone = phone or normalize_phone(sender)
        sender = None
    return jsonify({
        'asker_name': sender or '',
        'asker_phone': phone or '',
        'opened_at': opened_at or now(),
        'content': content,
        'raw_source_text': text,
    })


@app.route('/api/webhook/whatsapp', methods=['POST'])
def whatsapp_webhook():
    """Webhook תואם Twilio — יוצר שאלה חדשה מכל הודעה נכנסת. מוגן בטוקן."""
    if not webhook_token_ok():
        return jsonify({'error': 'forbidden'}), 403
    frm = request.form.get('From', '')
    body = (request.form.get('Body') or '').strip()
    profile = request.form.get('ProfileName', '')
    phone = normalize_phone(frm.replace('whatsapp:', ''))
    if not body and int(request.form.get('NumMedia') or 0) == 0:
        return _twiml()
    conn = get_db()
    # מניעת כפילות: אותה הודעה מאותו מספר ב-2 הדקות האחרונות
    recent = (now_dt() - timedelta(minutes=2)).isoformat(timespec='seconds')
    dup = conn.execute('SELECT id FROM questions WHERE asker_phone=? AND content=? AND opened_at>=?',
                       (phone, body, recent)).fetchone()
    if dup:
        conn.close()
        return _twiml()
    media_urls = []
    for i in range(int(request.form.get('NumMedia') or 0)):
        u = request.form.get(f'MediaUrl{i}')
        if u:
            media_urls.append(u)
    notes = ('קבצים מצורפים (וואטסאפ):\n' + '\n'.join(media_urls)) if media_urls else None
    raw = json.dumps({k: v for k, v in request.form.items()}, ensure_ascii=False)
    cur = conn.execute(
        'INSERT INTO questions (opened_at, asker_name, asker_phone, source, content, urgency, '
        'status, internal_notes, raw_source_text, updated_at) VALUES (?,?,?,?,?,?,?,?,?,?)',
        (now(), profile or phone, phone, 'וואטסאפ', body or '(הודעת מדיה ללא טקסט)',
         'רגיל', 'חדש', notes, raw, now()))
    log_history(conn, cur.lastrowid, 'created', 'התקבלה הודעת וואטסאפ', user_id=None)
    conn.commit()
    conn.close()
    return _twiml()


def _twiml():
    return Response('<?xml version="1.0" encoding="UTF-8"?><Response></Response>',
                    mimetype='application/xml')


# ---- שליחת וואטסאפ יוצא דרך Green API (מענה חוזר לפונה) ----
def _greenapi_creds():
    conn = get_db()
    iid = get_setting(conn, 'greenapi_instance_id')
    tok = get_setting(conn, 'greenapi_token')
    conn.close()
    return iid, tok


def wa_send_enabled():
    iid, tok = _greenapi_creds()
    return bool(iid and tok)


def send_whatsapp(phone, text):
    """שולח הודעת וואטסאפ למספר דרך Green API. זורק חריגה בכישלון."""
    iid, tok = _greenapi_creds()
    if not (iid and tok):
        raise RuntimeError('חיבור וואטסאפ לשליחה לא מוגדר (מסך הגדרות)')
    p = re.sub(r'\D', '', phone or '')
    if p.startswith('0'):
        p = '972' + p[1:]
    if not p:
        raise RuntimeError('לפונה אין מספר טלפון')
    import urllib.request
    import urllib.error
    req = urllib.request.Request(
        f'https://api.green-api.com/waInstance{iid}/sendMessage/{tok}',
        data=json.dumps({'chatId': p + '@c.us', 'message': text}).encode('utf-8'),
        headers={'Content-Type': 'application/json'}, method='POST')
    try:
        with urllib.request.urlopen(req, timeout=30) as resp:
            return json.loads(resp.read().decode('utf-8'))
    except urllib.error.HTTPError as e:
        raise RuntimeError(f'שירות הוואטסאפ דחה את הבקשה ({e.code}) — בדוק את פרטי החיבור בהגדרות')
    except urllib.error.URLError:
        raise RuntimeError('אין חיבור לשירות הוואטסאפ — נסה שוב')


@app.route('/api/questions/<int:qid>/send-answer', methods=['POST'])
def send_answer_whatsapp(qid):
    """שולח את המענה חזרה לפונה בוואטסאפ, מאותו מספר ייעודי."""
    conn = get_db()
    q = conn.execute('SELECT * FROM questions WHERE id=?', (qid,)).fetchone()
    conn.close()
    if not q:
        return jsonify({'error': 'שאלה לא נמצאה'}), 404
    if not q['asker_phone']:
        return jsonify({'error': 'לשאלה אין מספר טלפון של פונה'}), 400
    answer = q['approved_answer'] or (q['proposed_answer'] if not q['needs_approval'] else None)
    if not answer:
        return jsonify({'error': 'אין מענה מאושר לשליחה' if q['needs_approval'] else 'אין מענה לשליחה'}), 400
    try:
        send_whatsapp(q['asker_phone'], answer)
    except RuntimeError as e:
        return jsonify({'error': str(e)}), 502
    conn = get_db()
    conn.execute('UPDATE questions SET status=?, answered_at=COALESCE(answered_at, ?), updated_at=? WHERE id=?',
                 ('נענה', now(), now(), qid))
    log_history(conn, qid, 'answer_sent', f'המענה נשלח בוואטסאפ אל {q["asker_phone"]}')
    log_action(conn, 'send_answer_whatsapp', 'question', qid)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ---- Green API: בוט בקבוצה / מספר ייעודי (חיבור וואטסאפ דרך greenapi.com) ----
RELEVANCE_SCHEMA = {
    'type': 'object',
    'properties': {
        'relevant': {'type': 'boolean',
                     'description': 'האם ההודעה היא שאלה/פנייה שרלוונטית לתא ההסברה'},
        'topic': {'type': 'string', 'description': 'נושא מתאים מהרשימה, או מחרוזת ריקה'},
    },
    'required': ['relevant', 'topic'],
    'additionalProperties': False,
}


def _relevance_system():
    return (
        f'אתה מסנן הודעות עבור תא ההסברה של {DISTRICT_NAME} בפיקוד העורף. '
        'קבל הודעה מקבוצת וואטסאפ והחלט אם היא שאלה או פנייה שהתא צריך לטפל בה: '
        'שאלות על הנחיות התגוננות, לימודים, מקלטים, התקהלות, שירותים חיוניים, '
        'מידע לציבור, בקשות להבהרה או דיווח על פערי מידע. '
        'לא רלוונטי: שיחת חולין, ברכות, אישורי קבלה ("תודה", "קיבלתי", "👍"), '
        'תיאומים פנימיים והודעות מערכת. '
        f'אם רלוונטי — בחר נושא מהרשימה: {", ".join(SEED_TOPICS)}. '
        'החזר JSON בלבד.'
    )


def _classify_group_message(text):
    """מחזיר (relevant, topic). עם AI — סיווג חכם; בלעדיו — יוריסטיקה של סימן שאלה."""
    if ai_enabled():
        try:
            raw = call_claude(_relevance_system(), text[:2000], max_tokens=200,
                              schema=RELEVANCE_SCHEMA)
            data = json.loads(raw)
            return bool(data.get('relevant')), (data.get('topic') or None)
        except Exception:
            pass  # נפילה ליוריסטיקה — עדיף לקלוט מדי מאשר לפספס
    return ('?' in text or 'האם' in text or 'מתי' in text or 'איפה' in text), None


@app.route('/api/webhook/greenapi', methods=['POST'])
def greenapi_webhook():
    """Webhook ל-Green API: קולט הודעות ממספר ייעודי וגם מקבוצות שהבוט חבר בהן.
    הודעה פרטית — נקלטת תמיד; הודעת קבוצה — רק אם סווגה כרלוונטית לתא."""

    def wa_log(decision, detail=''):
        conn2 = get_db()
        log_action(conn2, 'wa_webhook', 'webhook', None,
                   f'{decision}' + (f' | {detail}' if detail else ''))
        conn2.commit()
        conn2.close()

    if not webhook_token_ok():
        # נרשם ביומן כדי שאפשר יהיה לאבחן כתובת webhook שהודבקה בלי טוקן / עם טוקן שגוי
        wa_log('❌ נדחה — טוקן שגוי או חסר בכתובת ה-webhook',
               'ודא שהכתובת ב-Green API כוללת את ?token=... במלואו')
        return jsonify({'error': 'forbidden'}), 403
    d = request.get_json(silent=True) or {}

    if d.get('typeWebhook') != 'incomingMessageReceived':
        wa_log('התעלמות — אירוע מסוג אחר', d.get('typeWebhook') or 'ללא סוג')
        return jsonify({'ok': True})
    md = d.get('messageData') or {}
    text = ((md.get('textMessageData') or {}).get('textMessage')
            or (md.get('extendedTextMessageData') or {}).get('text') or '').strip()
    sd = d.get('senderData') or {}
    chat_id = sd.get('chatId') or ''
    is_group = chat_id.endswith('@g.us')
    phone = normalize_phone((sd.get('sender') or '').split('@')[0])
    sender_name = sd.get('senderName') or phone
    chat_name = sd.get('chatName') or ''
    src_desc = f'קבוצה: {chat_name}' if is_group else f'פרטי: {sender_name}'
    if not text:
        wa_log('התעלמות — הודעה בלי טקסט', f'{src_desc} | {md.get("typeMessage") or ""}')
        return jsonify({'ok': True})
    topic = None
    if is_group:
        relevant, topic = _classify_group_message(text)
        if not relevant:
            wa_log('סונן — לא רלוונטי לתא', f'{src_desc} | {text[:60]}')
            return jsonify({'ok': True, 'skipped': 'not relevant'})
    conn = get_db()
    recent = (now_dt() - timedelta(minutes=2)).isoformat(timespec='seconds')
    dup = conn.execute('SELECT id FROM questions WHERE asker_phone=? AND content=? AND opened_at>=?',
                       (phone, text, recent)).fetchone()
    if dup:
        conn.close()
        wa_log('התעלמות — כפילות', f'{src_desc} | שאלה #{dup["id"]}')
        return jsonify({'ok': True, 'skipped': 'duplicate'})
    notes = f'נקלט מקבוצת וואטסאפ: {chat_name}' if is_group else None
    cur = conn.execute(
        'INSERT INTO questions (opened_at, asker_name, asker_phone, source, content, topic, '
        'urgency, status, internal_notes, raw_source_text, updated_at) '
        'VALUES (?,?,?,?,?,?,?,?,?,?,?)',
        (now(), sender_name, phone, 'וואטסאפ', text, topic, 'רגיל', 'חדש', notes,
         json.dumps(d, ensure_ascii=False)[:4000], now()))
    log_history(conn, cur.lastrowid, 'created',
                'התקבלה הודעת וואטסאפ' + (f' (קבוצה: {chat_name})' if is_group else ''),
                user_id=None)
    conn.commit()
    conn.close()
    wa_log(f'✅ נקלטה שאלה #{cur.lastrowid}', f'{src_desc} | {text[:60]}')
    return jsonify({'ok': True, 'question_id': cur.lastrowid})


# ---------------------------------------------------------------- מסמכים
def extract_text_from_file(path, ext):
    """חילוץ טקסט מ-txt/pdf/docx. מחזיר (טקסט, שגיאה)."""
    try:
        if ext == '.txt':
            with open(path, 'r', encoding='utf-8', errors='replace') as fh:
                return fh.read(), None
        if ext == '.pdf':
            from pypdf import PdfReader
            reader = PdfReader(path)
            return '\n'.join((p.extract_text() or '') for p in reader.pages), None
        if ext == '.docx':
            import docx
            document = docx.Document(path)
            parts = [p.text for p in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    parts.append(' | '.join(cell.text for cell in row.cells))
            return '\n'.join(parts), None
        return '', 'סוג קובץ לא נתמך לחילוץ טקסט'
    except Exception as e:
        return '', f'שגיאה בחילוץ טקסט: {e}'


@app.route('/api/documents')
def list_documents():
    conn = get_db()
    where, params = [], []
    if request.args.get('doc_type'):
        where.append('doc_type=?')
        params.append(request.args['doc_type'])
    q = request.args.get('q', '').strip()
    if q:
        for term in q.split():
            where.append('(title LIKE ? OR extracted_text LIKE ? OR ai_summary LIKE ?)')
            params.extend([f'%{term}%'] * 3)
    sql = 'SELECT id, title, doc_type, source, orig_name, size, ai_summary, insights_generated_at, ' \
          'uploaded_by_id, created_at FROM documents'
    if where:
        sql += ' WHERE ' + ' AND '.join(where)
    sql += ' ORDER BY created_at DESC LIMIT 300'
    conn2 = conn
    rows = []
    for r in conn.execute(sql, params):
        d = dict(r)
        d['uploaded_by_name'] = user_name(conn2, r['uploaded_by_id'])
        rows.append(d)
    conn.close()
    return jsonify(rows)


@app.route('/api/documents', methods=['POST'])
def upload_document():
    f = request.files.get('file')
    title = (request.form.get('title') or '').strip()
    if not f or not f.filename:
        return jsonify({'error': 'לא נבחר קובץ'}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in ('.txt', '.pdf', '.docx'):
        return jsonify({'error': 'ניתן להעלות רק קבצי PDF, Word (docx) או טקסט'}), 400
    stored = uuid.uuid4().hex + ext
    path = os.path.join(UPLOAD_DIR, stored)
    f.save(path)
    text, err = extract_text_from_file(path, ext)
    conn = get_db()
    cur = conn.execute(
        'INSERT INTO documents (title, doc_type, source, orig_name, stored_name, mime, size, '
        'extracted_text, uploaded_by_id, created_at) VALUES (?,?,?,?,?,?,?,?,?,?)',
        (title or f.filename, request.form.get('doc_type'), request.form.get('source'),
         f.filename, stored, f.mimetype, os.path.getsize(path), text, session['uid'], now()))
    doc_id = cur.lastrowid
    log_action(conn, 'upload_document', 'document', doc_id, f.filename)
    conn.commit()
    conn.close()
    warn = None
    if err:
        warn = err
    elif len((text or '').strip()) < 50:
        warn = 'החילוץ החזיר מעט מאוד טקסט — ייתכן שהקובץ סרוק. ניתן להדביק את הטקסט ידנית.'
    return jsonify({'id': doc_id, 'warning': warn})


@app.route('/api/documents/<int:doc_id>')
def get_document(doc_id):
    conn = get_db()
    r = conn.execute('SELECT * FROM documents WHERE id=?', (doc_id,)).fetchone()
    if not r:
        conn.close()
        return jsonify({'error': 'מסמך לא נמצא'}), 404
    d = dict(r)
    d['uploaded_by_name'] = user_name(conn, r['uploaded_by_id'])
    conn.close()
    return jsonify(d)


DOC_EDITABLE = ['title', 'doc_type', 'source', 'extracted_text', 'ai_summary', 'ai_key_points',
                'ai_messages', 'ai_qa', 'ai_gaps', 'ai_draft_message']


@app.route('/api/documents/<int:doc_id>', methods=['PUT'])
def update_document(doc_id):
    d = request.get_json(force=True)
    sets, params = [], []
    for f in DOC_EDITABLE:
        if f in d:
            sets.append(f'{f}=?')
            params.append(d[f])
    if sets:
        params.append(doc_id)
        conn = get_db()
        conn.execute(f'UPDATE documents SET {", ".join(sets)} WHERE id=?', params)
        log_action(conn, 'update_document', 'document', doc_id)
        conn.commit()
        conn.close()
    return jsonify({'ok': True})


@app.route('/api/documents/<int:doc_id>', methods=['DELETE'])
@roles_required('admin', 'lead')
def delete_document(doc_id):
    conn = get_db()
    r = conn.execute('SELECT stored_name FROM documents WHERE id=?', (doc_id,)).fetchone()
    if r and r['stored_name']:
        try:
            os.remove(os.path.join(UPLOAD_DIR, r['stored_name']))
        except OSError:
            pass
    conn.execute('DELETE FROM documents WHERE id=?', (doc_id,))
    log_action(conn, 'delete_document', 'document', doc_id)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/documents/<int:doc_id>/download')
def download_document(doc_id):
    conn = get_db()
    r = conn.execute('SELECT * FROM documents WHERE id=?', (doc_id,)).fetchone()
    conn.close()
    if not r or not r['stored_name']:
        return jsonify({'error': 'קובץ לא נמצא'}), 404
    return send_file(os.path.join(UPLOAD_DIR, r['stored_name']),
                     download_name=r['orig_name'], as_attachment=True)


@app.route('/api/documents/<int:doc_id>/insights', methods=['POST'])
def generate_insights(doc_id):
    if not ai_enabled():
        return jsonify({'ai': False, 'error': 'מפתח API לא מוגדר — ניתן למלא את השדות ידנית'}), 200
    conn = get_db()
    r = conn.execute('SELECT extracted_text, title FROM documents WHERE id=?', (doc_id,)).fetchone()
    conn.close()   # סוגרים לפני קריאת ה-API — לא מחזיקים נעילה
    if not r:
        return jsonify({'error': 'מסמך לא נמצא'}), 404
    text = (r['extracted_text'] or '').strip()
    if not text:
        return jsonify({'error': 'אין טקסט מחולץ במסמך. הדבק טקסט ידנית ונסה שוב.'}), 400
    try:
        raw = call_claude(INSIGHTS_SYSTEM,
                          f'שם המסמך: {r["title"]}\n\nתוכן המסמך:\n{text[:60000]}',
                          max_tokens=8000, schema=INSIGHTS_SCHEMA)
        data = json.loads(raw)
    except json.JSONDecodeError:
        return jsonify({'error': 'תשובת ה-AI לא תקינה — נסה שוב'}), 502
    except Exception as e:
        return ai_error(e)
    conn = get_db()
    conn.execute('UPDATE documents SET ai_summary=?, ai_key_points=?, ai_messages=?, ai_qa=?, '
                 'ai_gaps=?, ai_draft_message=?, insights_generated_at=?, insights_model=? WHERE id=?',
                 (data['summary'], data['key_points'], data['messages'], data['qa'],
                  data['gaps'], data['draft_message'], now(), get_model(), doc_id))
    log_action(conn, 'generate_insights', 'document', doc_id)
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'insights': data})


@app.route('/api/documents/<int:doc_id>/to-message', methods=['POST'])
def doc_to_message(doc_id):
    conn = get_db()
    r = conn.execute('SELECT * FROM documents WHERE id=?', (doc_id,)).fetchone()
    if not r:
        conn.close()
        return jsonify({'error': 'מסמך לא נמצא'}), 404
    body = r['ai_messages'] or r['ai_summary'] or ''
    cur = conn.execute('INSERT INTO messages (title, body, status, source_document_id, '
                       'created_by_id, created_at, updated_at) VALUES (?,?,?,?,?,?,?)',
                       (f'מסר מתוך: {r["title"]}', body, 'טיוטה', doc_id, session['uid'], now(), now()))
    log_action(conn, 'doc_to_message', 'message', cur.lastrowid)
    conn.commit()
    conn.close()
    return jsonify({'id': cur.lastrowid})


# ---------------------------------------------------------------- מסרים
@app.route('/api/messages')
def list_messages():
    conn = get_db()
    where, params = [], []
    if request.args.get('status'):
        where.append('status=?')
        params.append(request.args['status'])
    sql = 'SELECT * FROM messages'
    if where:
        sql += ' WHERE ' + ' AND '.join(where)
    sql += ' ORDER BY updated_at DESC, id DESC LIMIT 300'
    rows = []
    for r in conn.execute(sql, params):
        d = dict(r)
        d['created_by_name'] = user_name(conn, r['created_by_id'])
        d['approved_by_name'] = user_name(conn, r['approved_by_id'])
        rows.append(d)
    conn.close()
    return jsonify(rows)


@app.route('/api/messages', methods=['POST'])
def create_message():
    d = request.get_json(force=True)
    if not (d.get('title') or '').strip() or not (d.get('body') or '').strip():
        return jsonify({'error': 'נושא ונוסח המסר חובה'}), 400
    conn = get_db()
    cur = conn.execute(
        'INSERT INTO messages (title, body, status, audience, valid_from, valid_until, '
        'source_document_id, notes, created_by_id, created_at, updated_at) '
        'VALUES (?,?,?,?,?,?,?,?,?,?,?)',
        (d['title'].strip(), d['body'].strip(), d.get('status') or 'טיוטה', d.get('audience'),
         d.get('valid_from'), d.get('valid_until'), d.get('source_document_id'),
         d.get('notes'), session['uid'], now(), now()))
    log_action(conn, 'create_message', 'message', cur.lastrowid)
    conn.commit()
    conn.close()
    return jsonify({'id': cur.lastrowid})


@app.route('/api/messages/<int:mid>', methods=['PUT'])
def update_message(mid):
    d = request.get_json(force=True)
    sets, params = [], []
    for f in ('title', 'body', 'status', 'audience', 'valid_from', 'valid_until', 'notes'):
        if f in d:
            if f == 'status' and d[f] not in MSG_STATUSES:
                return jsonify({'error': 'סטטוס לא חוקי'}), 400
            sets.append(f'{f}=?')
            params.append(d[f])
    if sets:
        sets.append('updated_at=?')
        params.append(now())
        params.append(mid)
        conn = get_db()
        conn.execute(f'UPDATE messages SET {", ".join(sets)} WHERE id=?', params)
        log_action(conn, 'update_message', 'message', mid)
        conn.commit()
        conn.close()
    return jsonify({'ok': True})


@app.route('/api/messages/<int:mid>/approve', methods=['POST'])
@roles_required('admin', 'lead')
def approve_message(mid):
    conn = get_db()
    conn.execute('UPDATE messages SET status=?, approved_by_id=?, updated_at=? WHERE id=?',
                 ('פעיל', session['uid'], now(), mid))
    log_action(conn, 'approve_message', 'message', mid)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/messages/<int:mid>/replace', methods=['POST'])
def replace_message(mid):
    """יוצר גרסה חדשה של מסר; הישן מסומן 'הוחלף'."""
    d = request.get_json(force=True)
    conn = get_db()
    old = conn.execute('SELECT * FROM messages WHERE id=?', (mid,)).fetchone()
    if not old:
        conn.close()
        return jsonify({'error': 'מסר לא נמצא'}), 404
    cur = conn.execute(
        'INSERT INTO messages (title, body, status, audience, valid_from, valid_until, '
        'source_document_id, notes, created_by_id, created_at, updated_at) '
        'VALUES (?,?,?,?,?,?,?,?,?,?,?)',
        (d.get('title') or old['title'], d.get('body') or old['body'], 'טיוטה',
         d.get('audience') or old['audience'], d.get('valid_from'), d.get('valid_until'),
         old['source_document_id'], old['notes'], session['uid'], now(), now()))
    new_id = cur.lastrowid
    conn.execute('UPDATE messages SET status=?, replaced_by_id=?, updated_at=? WHERE id=?',
                 ('הוחלף', new_id, now(), mid))
    log_action(conn, 'replace_message', 'message', mid, f'הוחלף ע"י #{new_id}')
    conn.commit()
    conn.close()
    return jsonify({'id': new_id})


# ---------------------------------------------------------------- פעילות
@app.route('/api/activities')
def list_activities():
    conn = get_db()
    where, params = [], []
    for field in ('activity_type', 'status', 'topic', 'authority'):
        if request.args.get(field):
            where.append(f'{field}=?')
            params.append(request.args[field])
    if request.args.get('performer'):
        where.append('performed_by_id=?')
        params.append(request.args['performer'])
    if request.args.get('from'):
        where.append('performed_at >= ?')
        params.append(request.args['from'])
    if request.args.get('to'):
        where.append('performed_at <= ?')
        params.append(request.args['to'] + 'T23:59:59')
    sql = 'SELECT * FROM activities'
    if where:
        sql += ' WHERE ' + ' AND '.join(where)
    sql += ' ORDER BY performed_at DESC LIMIT 500'
    rows = []
    for r in conn.execute(sql, params):
        d = dict(r)
        d['performed_by_name'] = user_name(conn, r['performed_by_id'])
        rows.append(d)
    conn.close()
    return jsonify(rows)


@app.route('/api/activities', methods=['POST'])
def create_activity():
    d = request.get_json(force=True)
    if not (d.get('description') or '').strip():
        return jsonify({'error': 'תוכן הפעולה חובה'}), 400
    conn = get_db()
    cur = conn.execute(
        'INSERT INTO activities (activity_type, description, topic, authority, audience, status, '
        'performed_at, performed_by_id, question_id, document_id, message_id, notes, created_at) '
        'VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?)',
        (d.get('activity_type') or 'אחר', d['description'].strip(), d.get('topic'),
         d.get('authority'), d.get('audience'), d.get('status') or 'בוצע',
         d.get('performed_at') or now(), session['uid'], d.get('question_id'),
         d.get('document_id'), d.get('message_id'), d.get('notes'), now()))
    log_action(conn, 'create_activity', 'activity', cur.lastrowid)
    conn.commit()
    conn.close()
    return jsonify({'id': cur.lastrowid})


@app.route('/api/activities/<int:aid>', methods=['PUT'])
def update_activity(aid):
    d = request.get_json(force=True)
    sets, params = [], []
    for f in ('activity_type', 'description', 'topic', 'authority', 'audience', 'status',
              'performed_at', 'notes'):
        if f in d:
            sets.append(f'{f}=?')
            params.append(d[f])
    if sets:
        params.append(aid)
        conn = get_db()
        conn.execute(f'UPDATE activities SET {", ".join(sets)} WHERE id=?', params)
        conn.commit()
        conn.close()
    return jsonify({'ok': True})


@app.route('/api/activities/<int:aid>', methods=['DELETE'])
@roles_required('admin', 'lead')
def delete_activity(aid):
    conn = get_db()
    conn.execute('DELETE FROM activities WHERE id=?', (aid,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ---------------------------------------------------------------- הודעות להפצה
DRAFT_SYSTEM = (
    f'אתה קצין הסברה ב{DISTRICT_NAME} של פיקוד העורף. נסח הודעה קצרה, ברורה ומעשית להפצה, '
    'מותאמת לקהל היעד שצוין. כתוב בעברית פשוטה. החזר את נוסח ההודעה בלבד, '
    'ללא הקדמות והסברים.'
)


def _outgoing_source_text(conn, source_type, source_id):
    if source_type == 'question' and source_id:
        r = conn.execute('SELECT content, approved_answer, proposed_answer FROM questions WHERE id=?',
                         (source_id,)).fetchone()
        if r:
            return f'שאלה: {r["content"]}\nמענה: {r["approved_answer"] or r["proposed_answer"] or ""}'
    if source_type == 'document' and source_id:
        r = conn.execute('SELECT title, ai_summary, ai_draft_message, extracted_text FROM documents '
                         'WHERE id=?', (source_id,)).fetchone()
        if r:
            return r['ai_draft_message'] or r['ai_summary'] or (r['extracted_text'] or '')[:4000]
    if source_type == 'message' and source_id:
        r = conn.execute('SELECT body FROM messages WHERE id=?', (source_id,)).fetchone()
        if r:
            return r['body']
    return ''


@app.route('/api/outgoing')
def list_outgoing():
    conn = get_db()
    rows = []
    for r in conn.execute('SELECT * FROM outgoing_messages ORDER BY created_at DESC LIMIT 300'):
        d = dict(r)
        d['created_by_name'] = user_name(conn, r['created_by_id'])
        d['approved_by_name'] = user_name(conn, r['approved_by_id'])
        rows.append(d)
    conn.close()
    return jsonify(rows)


@app.route('/api/outgoing/draft', methods=['POST'])
def draft_outgoing():
    """מחזיר טיוטת הודעה — עם AI אם מוגדר, אחרת טקסט המקור."""
    d = request.get_json(force=True)
    conn = get_db()
    src = _outgoing_source_text(conn, d.get('source_type'), d.get('source_id'))
    conn.close()
    manual = (d.get('manual_text') or '').strip()
    base = src or manual
    if not base:
        return jsonify({'error': 'אין מקור מידע להודעה'}), 400
    audience = d.get('audience') or 'ציבור'
    if not ai_enabled():
        return jsonify({'body': base, 'ai': False})
    try:
        body = call_claude(DRAFT_SYSTEM,
                           f'קהל היעד: {audience}\n\nהמידע שממנו יש לנסח את ההודעה:\n{base[:20000]}',
                           max_tokens=2000)
        return jsonify({'body': body.strip(), 'ai': True})
    except Exception as e:
        return ai_error(e)


@app.route('/api/outgoing', methods=['POST'])
def create_outgoing():
    d = request.get_json(force=True)
    if not (d.get('body') or '').strip():
        return jsonify({'error': 'נוסח ההודעה חובה'}), 400
    conn = get_db()
    cur = conn.execute(
        'INSERT INTO outgoing_messages (source_type, source_id, audience, body, status, '
        'needs_approval, created_by_id, ai_generated, created_at) VALUES (?,?,?,?,?,?,?,?,?)',
        (d.get('source_type') or 'manual', d.get('source_id'), d.get('audience'),
         d['body'].strip(), 'טיוטה', 1 if d.get('needs_approval') else 0,
         session['uid'], 1 if d.get('ai_generated') else 0, now()))
    log_action(conn, 'create_outgoing', 'outgoing', cur.lastrowid)
    conn.commit()
    conn.close()
    return jsonify({'id': cur.lastrowid})


@app.route('/api/outgoing/<int:oid>', methods=['PUT'])
def update_outgoing(oid):
    d = request.get_json(force=True)
    sets, params = [], []
    for f in ('body', 'audience', 'needs_approval'):
        if f in d:
            sets.append(f'{f}=?')
            params.append(1 if (f == 'needs_approval' and d[f]) else (0 if f == 'needs_approval' else d[f]))
    if sets:
        params.append(oid)
        conn = get_db()
        conn.execute(f'UPDATE outgoing_messages SET {", ".join(sets)} WHERE id=?', params)
        conn.commit()
        conn.close()
    return jsonify({'ok': True})


@app.route('/api/outgoing/<int:oid>/approve', methods=['POST'])
@roles_required('admin', 'lead')
def approve_outgoing(oid):
    conn = get_db()
    conn.execute('UPDATE outgoing_messages SET status=?, approved_by_id=? WHERE id=?',
                 ('מאושר', session['uid'], oid))
    log_action(conn, 'approve_outgoing', 'outgoing', oid)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/outgoing/<int:oid>/distribute', methods=['POST'])
def distribute_outgoing(oid):
    conn = get_db()
    o = conn.execute('SELECT * FROM outgoing_messages WHERE id=?', (oid,)).fetchone()
    if not o:
        conn.close()
        return jsonify({'error': 'הודעה לא נמצאה'}), 404
    if o['needs_approval'] and o['status'] != 'מאושר' and session.get('role') not in ('admin', 'lead'):
        conn.close()
        return jsonify({'error': 'ההודעה דורשת אישור לפני הפצה'}), 403
    cur = conn.execute(
        'INSERT INTO activities (activity_type, description, audience, status, performed_at, '
        'performed_by_id, question_id, document_id, message_id, created_at) '
        'VALUES (?,?,?,?,?,?,?,?,?,?)',
        ('הודעה שהופצה', o['body'], o['audience'], 'הופץ', now(), session['uid'],
         o['source_id'] if o['source_type'] == 'question' else None,
         o['source_id'] if o['source_type'] == 'document' else None,
         o['source_id'] if o['source_type'] == 'message' else None, now()))
    conn.execute('UPDATE outgoing_messages SET status=?, distributed_at=?, activity_id=? WHERE id=?',
                 ('הופץ', now(), cur.lastrowid, oid))
    log_action(conn, 'distribute_outgoing', 'outgoing', oid)
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'activity_id': cur.lastrowid})


# ---------------------------------------------------------------- בנק הודעות מוכנות
@app.route('/api/canned')
def list_canned():
    conn = get_db()
    where, params = ['active=1'], []
    if request.args.get('category'):
        where.append('category=?')
        params.append(request.args['category'])
    rows = []
    for r in conn.execute('SELECT * FROM canned_messages WHERE ' + ' AND '.join(where) +
                          ' ORDER BY category, title LIMIT 300', params):
        d = dict(r)
        d['created_by_name'] = user_name(conn, r['created_by_id'])
        rows.append(d)
    conn.close()
    return jsonify(rows)


@app.route('/api/canned', methods=['POST'])
def create_canned():
    d = request.get_json(force=True)
    if not (d.get('title') or '').strip() or not (d.get('body') or '').strip():
        return jsonify({'error': 'כותרת ונוסח חובה'}), 400
    conn = get_db()
    cur = conn.execute(
        'INSERT INTO canned_messages (title, category, body, audience, notes, active, '
        'created_by_id, created_at, updated_at) VALUES (?,?,?,?,?,1,?,?,?)',
        (d['title'].strip(), d.get('category'), d['body'].strip(), d.get('audience'),
         d.get('notes'), session['uid'], now(), now()))
    log_action(conn, 'create_canned', 'canned', cur.lastrowid)
    conn.commit()
    conn.close()
    return jsonify({'id': cur.lastrowid})


@app.route('/api/canned/<int:cid>', methods=['PUT'])
def update_canned(cid):
    d = request.get_json(force=True)
    sets, params = [], []
    for f in ('title', 'category', 'body', 'audience', 'notes'):
        if f in d:
            sets.append(f'{f}=?')
            params.append(d[f])
    if sets:
        sets.append('updated_at=?')
        params.append(now())
        params.append(cid)
        conn = get_db()
        conn.execute(f'UPDATE canned_messages SET {", ".join(sets)} WHERE id=?', params)
        log_action(conn, 'update_canned', 'canned', cid)
        conn.commit()
        conn.close()
    return jsonify({'ok': True})


@app.route('/api/canned/<int:cid>', methods=['DELETE'])
@roles_required('admin', 'lead')
def delete_canned(cid):
    conn = get_db()
    conn.execute('UPDATE canned_messages SET active=0 WHERE id=?', (cid,))
    log_action(conn, 'delete_canned', 'canned', cid)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ---------------------------------------------------------------- גלריית חומרי הסברה
@app.route('/api/materials')
def list_materials():
    conn = get_db()
    where, params = [], []
    if request.args.get('category'):
        where.append('category=?')
        params.append(request.args['category'])
    sql = 'SELECT * FROM materials'
    if where:
        sql += ' WHERE ' + ' AND '.join(where)
    sql += ' ORDER BY created_at DESC LIMIT 300'
    rows = []
    for r in conn.execute(sql, params):
        d = dict(r)
        d['uploaded_by_name'] = user_name(conn, r['uploaded_by_id'])
        d['is_image'] = (r['mime'] or '').startswith('image/')
        rows.append(d)
    conn.close()
    return jsonify(rows)


@app.route('/api/materials', methods=['POST'])
def upload_material():
    f = request.files.get('file')
    if not f or not f.filename:
        return jsonify({'error': 'לא נבחר קובץ'}), 400
    ext = os.path.splitext(f.filename)[1].lower()
    if ext not in MATERIAL_EXTS:
        return jsonify({'error': 'סוג קובץ לא נתמך. מותר: תמונות, PDF, וידאו (mp4), מצגות ומסמכים'}), 400
    stored = uuid.uuid4().hex + ext
    path = os.path.join(UPLOAD_DIR, stored)
    f.save(path)
    conn = get_db()
    cur = conn.execute(
        'INSERT INTO materials (title, category, description, orig_name, stored_name, mime, '
        'size, uploaded_by_id, created_at) VALUES (?,?,?,?,?,?,?,?,?)',
        ((request.form.get('title') or f.filename).strip(), request.form.get('category'),
         request.form.get('description'), f.filename, stored, f.mimetype,
         os.path.getsize(path), session['uid'], now()))
    log_action(conn, 'upload_material', 'material', cur.lastrowid, f.filename)
    conn.commit()
    conn.close()
    return jsonify({'id': cur.lastrowid})


@app.route('/api/materials/<int:mid>/file')
def material_file(mid):
    conn = get_db()
    r = conn.execute('SELECT * FROM materials WHERE id=?', (mid,)).fetchone()
    conn.close()
    if not r:
        return jsonify({'error': 'קובץ לא נמצא'}), 404
    as_attach = request.args.get('dl') == '1'
    return send_file(os.path.join(UPLOAD_DIR, r['stored_name']),
                     download_name=r['orig_name'], as_attachment=as_attach)


@app.route('/api/materials/<int:mid>', methods=['PUT'])
def update_material(mid):
    d = request.get_json(force=True)
    sets, params = [], []
    for f in ('title', 'category', 'description'):
        if f in d:
            sets.append(f'{f}=?')
            params.append(d[f])
    if sets:
        params.append(mid)
        conn = get_db()
        conn.execute(f'UPDATE materials SET {", ".join(sets)} WHERE id=?', params)
        conn.commit()
        conn.close()
    return jsonify({'ok': True})


@app.route('/api/materials/<int:mid>', methods=['DELETE'])
@roles_required('admin', 'lead')
def delete_material(mid):
    conn = get_db()
    r = conn.execute('SELECT stored_name FROM materials WHERE id=?', (mid,)).fetchone()
    if r and r['stored_name']:
        try:
            os.remove(os.path.join(UPLOAD_DIR, r['stored_name']))
        except OSError:
            pass
    conn.execute('DELETE FROM materials WHERE id=?', (mid,))
    log_action(conn, 'delete_material', 'material', mid)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


# ---------------------------------------------------------------- סיכום משמרת
def gather_stats(conn, start, end):
    end_full = end if 'T' in end else end + 'T23:59:59'
    args = (start, end_full)
    stats = {}
    stats['opened'] = [dict(r) for r in conn.execute(
        'SELECT id, content, topic, urgency, status FROM questions WHERE opened_at BETWEEN ? AND ? '
        'ORDER BY opened_at', args)]
    stats['closed'] = [dict(r) for r in conn.execute(
        'SELECT id, content, topic FROM questions WHERE closed_at BETWEEN ? AND ?', args)]
    stats['open_now'] = [dict(r) for r in conn.execute(
        'SELECT id, content, topic, urgency, status FROM questions WHERE status IN (%s) '
        'ORDER BY CASE urgency WHEN \'מיידי\' THEN 0 WHEN \'דחוף\' THEN 1 ELSE 2 END'
        % ','.join('?' * len(OPEN_STATUSES)), OPEN_STATUSES)]
    stats['urgent_open'] = [q for q in stats['open_now'] if q['urgency'] in ('דחוף', 'מיידי')]
    stats['awaiting_approval'] = [q for q in stats['open_now'] if q['status'] == 'ממתין לאישור']
    stats['activities'] = [dict(r) for r in conn.execute(
        'SELECT id, activity_type, description, audience FROM activities '
        'WHERE performed_at BETWEEN ? AND ? ORDER BY performed_at', args)]
    stats['distributed'] = [a for a in stats['activities'] if a['activity_type'] == 'הודעה שהופצה']
    stats['documents'] = [dict(r) for r in conn.execute(
        'SELECT id, title, doc_type FROM documents WHERE created_at BETWEEN ? AND ?', args)]
    stats['messages_updated'] = [dict(r) for r in conn.execute(
        'SELECT id, title, status FROM messages WHERE updated_at BETWEEN ? AND ?', args)]
    stats['gaps'] = [dict(r) for r in conn.execute(
        "SELECT id, title, ai_gaps FROM documents WHERE ai_gaps IS NOT NULL AND ai_gaps != ''")]
    return stats


def build_summary_text(stats, start, end):
    def sec(title, items, fmt):
        lines = [f'\n## {title} ({len(items)})']
        if not items:
            lines.append('- אין')
        for it in items:
            lines.append('- ' + fmt(it))
        return lines
    L = [f'# סיכום משמרת — תא הסברה {DISTRICT_NAME}', f'תקופה: {start} עד {end}', f'הופק: {now()}']
    L += sec('שאלות שנפתחו', stats['opened'],
             lambda q: f'[#{q["id"]}] {q["content"][:90]} ({q["topic"] or "ללא נושא"}, {q["status"]})')
    L += sec('שאלות שנסגרו', stats['closed'], lambda q: f'[#{q["id"]}] {q["content"][:90]}')
    L += sec('שאלות פתוחות כעת', stats['open_now'],
             lambda q: f'[#{q["id"]}] {q["content"][:90]} ({q["urgency"]}, {q["status"]})')
    L += sec('דחוף להמשך', stats['urgent_open'],
             lambda q: f'[#{q["id"]}] {q["content"][:90]} ({q["urgency"]})')
    L += sec('ממתין לאישור', stats['awaiting_approval'], lambda q: f'[#{q["id"]}] {q["content"][:90]}')
    L += sec('פעולות שבוצעו', stats['activities'],
             lambda a: f'{a["activity_type"]}: {a["description"][:90]}')
    L += sec('הודעות שהופצו', stats['distributed'],
             lambda a: f'{a["description"][:90]} (ל{a["audience"] or "קהל כללי"})')
    L += sec('מסמכים שהועלו', stats['documents'], lambda d: f'{d["title"]} ({d["doc_type"] or "-"})')
    L += sec('מסרים שעודכנו', stats['messages_updated'], lambda m: f'{m["title"]} ({m["status"]})')
    L += sec('פערים פתוחים', stats['gaps'], lambda g: f'{g["title"]}: {(g["ai_gaps"] or "")[:120]}')
    return '\n'.join(L)


SUMMARY_SYSTEM = (
    f'אתה קצין הסברה ב{DISTRICT_NAME} של פיקוד העורף. קבל נתוני משמרת מובנים וכתוב סיכום משמרת '
    'נרטיבי, ברור ותמציתי בעברית, במבנה הקבוע הבא (כותרות):\n'
    '1. מה קרה במשמרת\n2. מה הופץ\n3. מה טופל\n4. מה עדיין פתוח\n5. מה דחוף להמשך\n'
    '6. מה דורש אישור\n7. המלצות למשמרת הבאה\n'
    'הסתמך אך ורק על הנתונים שסופקו. אם סעיף ריק — כתוב "אין".'
)


@app.route('/api/summaries')
def list_summaries():
    conn = get_db()
    rows = []
    for r in conn.execute('SELECT * FROM shift_summaries ORDER BY created_at DESC LIMIT 100'):
        d = dict(r)
        d['created_by_name'] = user_name(conn, r['created_by_id'])
        rows.append(d)
    conn.close()
    return jsonify(rows)


@app.route('/api/summaries/preview')
def preview_summary():
    conn = get_db()
    start = request.args.get('from')
    if not start:
        last = conn.execute('SELECT MAX(period_end) me FROM shift_summaries').fetchone()
        start = last['me'] or (today() + 'T00:00:00')
    end = request.args.get('to') or now()
    stats = gather_stats(conn, start, end)
    conn.close()
    body = build_summary_text(stats, start, end)
    counts = {k: len(v) for k, v in stats.items()}
    return jsonify({'period_start': start, 'period_end': end, 'body': body, 'counts': counts})


@app.route('/api/summaries/draft', methods=['POST'])
def draft_summary():
    """ניסוח נרטיבי עם AI על בסיס נתוני התקופה."""
    d = request.get_json(force=True)
    start, end = d.get('from'), d.get('to') or now()
    if not start:
        return jsonify({'error': 'חסר תאריך התחלה'}), 400
    conn = get_db()
    stats = gather_stats(conn, start, end)
    conn.close()
    template = build_summary_text(stats, start, end)
    if not ai_enabled():
        return jsonify({'body': template, 'ai': False})
    try:
        body = call_claude(SUMMARY_SYSTEM, template[:40000], max_tokens=4000)
        return jsonify({'body': body.strip(), 'ai': True})
    except Exception as e:
        return ai_error(e)


@app.route('/api/summaries', methods=['POST'])
def save_summary():
    d = request.get_json(force=True)
    if not (d.get('body') or '').strip():
        return jsonify({'error': 'תוכן הסיכום ריק'}), 400
    conn = get_db()
    stats = gather_stats(conn, d.get('period_start') or today(), d.get('period_end') or now())
    counts = {k: len(v) for k, v in stats.items()}
    cur = conn.execute(
        'INSERT INTO shift_summaries (period_start, period_end, body, stats_json, ai_generated, '
        'created_by_id, created_at) VALUES (?,?,?,?,?,?,?)',
        (d.get('period_start') or today(), d.get('period_end') or now(), d['body'].strip(),
         json.dumps(counts, ensure_ascii=False), 1 if d.get('ai_generated') else 0,
         session['uid'], now()))
    log_action(conn, 'save_summary', 'summary', cur.lastrowid)
    conn.commit()
    conn.close()
    return jsonify({'id': cur.lastrowid})


# ---------------------------------------------------------------- דשבורד ודוחות
@app.route('/api/dashboard')
def dashboard():
    conn = get_db()
    t0 = today() + 'T00:00:00'
    d = {}
    d['open'] = conn.execute("SELECT COUNT(*) c FROM questions WHERE status IN ('חדש','דורש המשך טיפול')").fetchone()['c']
    d['in_progress'] = conn.execute("SELECT COUNT(*) c FROM questions WHERE status IN ('בטיפול','ממתין למידע')").fetchone()['c']
    d['awaiting_approval'] = conn.execute("SELECT COUNT(*) c FROM questions WHERE status='ממתין לאישור'").fetchone()['c']
    d['closed_today'] = conn.execute("SELECT COUNT(*) c FROM questions WHERE closed_at>=?", (t0,)).fetchone()['c']
    d['activities_today'] = conn.execute('SELECT COUNT(*) c FROM activities WHERE performed_at>=?', (t0,)).fetchone()['c']
    d['distributed_today'] = conn.execute(
        "SELECT COUNT(*) c FROM activities WHERE performed_at>=? AND activity_type='הודעה שהופצה'", (t0,)).fetchone()['c']
    d['urgent_open'] = [dict(r) for r in conn.execute(
        "SELECT id, content, urgency, status, topic FROM questions WHERE status IN (%s) "
        "AND urgency IN ('דחוף','מיידי') ORDER BY CASE urgency WHEN 'מיידי' THEN 0 ELSE 1 END, opened_at"
        % ','.join('?' * len(OPEN_STATUSES)), OPEN_STATUSES)]
    d['recent_documents'] = [dict(r) for r in conn.execute(
        'SELECT id, title, doc_type, created_at FROM documents ORDER BY created_at DESC LIMIT 5')]
    d['active_messages'] = [dict(r) for r in conn.execute(
        "SELECT id, title, audience, updated_at FROM messages WHERE status='פעיל' "
        'ORDER BY updated_at DESC LIMIT 8')]
    d['recurring_topics'] = [dict(r) for r in conn.execute(
        'SELECT topic, COUNT(*) c FROM questions WHERE topic IS NOT NULL AND topic != "" '
        'GROUP BY topic HAVING c >= 2 ORDER BY c DESC LIMIT 8')]
    d['open_gaps'] = [dict(r) for r in conn.execute(
        "SELECT id, title, ai_gaps FROM documents WHERE ai_gaps IS NOT NULL AND ai_gaps != '' "
        'ORDER BY created_at DESC LIMIT 5')]
    conn.close()
    return jsonify(d)


@app.route('/api/reports/<kind>')
def report(kind):
    conn = get_db()
    title, lines = '', []
    if kind == 'status_summary':
        title = 'סיכום מצב הסברתי'
        t0 = today() + 'T00:00:00'
        stats = gather_stats(conn, t0, now())
        lines = build_summary_text(stats, t0, now()).splitlines()
        lines[0] = f'# סיכום מצב הסברתי — תא הסברה {DISTRICT_NAME}'
    elif kind == 'open_questions':
        title = 'רשימת שאלות פתוחות'
        lines = [f'# {title}', f'הופק: {now()}', '']
        for q in conn.execute(
                'SELECT * FROM questions WHERE status IN (%s) ORDER BY '
                "CASE urgency WHEN 'מיידי' THEN 0 WHEN 'דחוף' THEN 1 ELSE 2 END, opened_at"
                % ','.join('?' * len(OPEN_STATUSES)), OPEN_STATUSES):
            lines.append(f'[#{q["id"]}] ({q["urgency"]}, {q["status"]}) {q["content"]}')
            lines.append(f'    נושא: {q["topic"] or "-"} | רשות: {q["authority"] or "-"} | '
                         f'פונה: {q["asker_name"] or "-"} | נפתחה: {q["opened_at"]}')
            lines.append('')
    elif kind == 'actions':
        title = 'רשימת פעולות שבוצעו היום'
        t0 = today() + 'T00:00:00'
        lines = [f'# {title}', f'הופק: {now()}', '']
        for a in conn.execute('SELECT * FROM activities WHERE performed_at>=? ORDER BY performed_at', (t0,)):
            lines.append(f'- {a["performed_at"][11:16]} | {a["activity_type"]} | {a["description"]}'
                         + (f' (ל{a["audience"]})' if a['audience'] else ''))
    elif kind == 'messages_doc':
        title = 'מסמך מסרים עדכני'
        lines = [f'# {title}', f'הופק: {now()}', '']
        for m in conn.execute("SELECT * FROM messages WHERE status='פעיל' ORDER BY updated_at DESC"):
            lines.append(f'## {m["title"]}' + (f' (קהל יעד: {m["audience"]})' if m['audience'] else ''))
            lines.append(m['body'])
            if m['valid_until']:
                lines.append(f'תוקף עד: {m["valid_until"]}')
            lines.append('')
    elif kind == 'gaps':
        title = 'רשימת פערים פתוחים'
        lines = [f'# {title}', f'הופק: {now()}', '']
        for g in conn.execute("SELECT title, ai_gaps FROM documents WHERE ai_gaps IS NOT NULL AND ai_gaps != ''"):
            lines.append(f'## מתוך: {g["title"]}')
            lines.append(g['ai_gaps'])
            lines.append('')
        for q in conn.execute("SELECT id, content FROM questions WHERE status='ממתין למידע'"):
            lines.append(f'- שאלה ממתינה למידע [#{q["id"]}]: {q["content"][:120]}')
    else:
        conn.close()
        return jsonify({'error': 'דוח לא מוכר'}), 404
    conn.close()
    return jsonify({'title': title, 'text': '\n'.join(lines)})


# ---------------------------------------------------------------- חיפוש
@app.route('/api/search')
def search():
    q = request.args.get('q', '').strip()
    if not q:
        return jsonify([])
    types = (request.args.get('types') or
             'questions,documents,messages,activities,summaries,canned,materials').split(',')
    terms = q.split()

    def like_clause(fields):
        parts, params = [], []
        for term in terms:
            sub = ' OR '.join(f'{f} LIKE ?' for f in fields)
            parts.append(f'({sub})')
            params.extend([f'%{term}%'] * len(fields))
        return ' AND '.join(parts), params

    conn = get_db()
    results = []
    if 'questions' in types:
        w, p = like_clause(['content', 'proposed_answer', 'approved_answer', 'asker_name', 'internal_notes'])
        for r in conn.execute(f'SELECT id, content, status, topic, opened_at FROM questions WHERE {w} '
                              'ORDER BY opened_at DESC LIMIT 50', p):
            results.append({'type': 'question', 'id': r['id'], 'title': r['content'][:120],
                            'meta': f'{r["status"]} · {r["topic"] or ""} · {r["opened_at"][:10]}'})
    if 'documents' in types:
        w, p = like_clause(['title', 'extracted_text', 'ai_summary', 'ai_messages', 'ai_qa'])
        for r in conn.execute(f'SELECT id, title, doc_type, created_at FROM documents WHERE {w} '
                              'ORDER BY created_at DESC LIMIT 50', p):
            results.append({'type': 'document', 'id': r['id'], 'title': r['title'],
                            'meta': f'{r["doc_type"] or "מסמך"} · {r["created_at"][:10]}'})
    if 'messages' in types:
        w, p = like_clause(['title', 'body'])
        for r in conn.execute(f'SELECT id, title, status, updated_at FROM messages WHERE {w} '
                              'ORDER BY updated_at DESC LIMIT 50', p):
            results.append({'type': 'message', 'id': r['id'], 'title': r['title'],
                            'meta': f'{r["status"]} · {(r["updated_at"] or "")[:10]}'})
    if 'activities' in types:
        w, p = like_clause(['description', 'notes'])
        for r in conn.execute(f'SELECT id, activity_type, description, performed_at FROM activities '
                              f'WHERE {w} ORDER BY performed_at DESC LIMIT 50', p):
            results.append({'type': 'activity', 'id': r['id'], 'title': r['description'][:120],
                            'meta': f'{r["activity_type"]} · {r["performed_at"][:10]}'})
    if 'summaries' in types:
        w, p = like_clause(['body'])
        for r in conn.execute(f'SELECT id, period_start, period_end, created_at FROM shift_summaries '
                              f'WHERE {w} ORDER BY created_at DESC LIMIT 20', p):
            results.append({'type': 'summary', 'id': r['id'],
                            'title': f'סיכום משמרת {r["period_start"][:10]} — {r["period_end"][:10]}',
                            'meta': r['created_at'][:16]})
    if 'canned' in types:
        w, p = like_clause(['title', 'body'])
        for r in conn.execute(f'SELECT id, title, category FROM canned_messages WHERE active=1 AND ({w}) '
                              'ORDER BY title LIMIT 30', p):
            results.append({'type': 'canned', 'id': r['id'], 'title': r['title'],
                            'meta': r['category'] or 'הודעה מוכנה'})
    if 'materials' in types:
        w, p = like_clause(['title', 'description', 'orig_name'])
        for r in conn.execute(f'SELECT id, title, category, created_at FROM materials WHERE {w} '
                              'ORDER BY created_at DESC LIMIT 30', p):
            results.append({'type': 'material', 'id': r['id'], 'title': r['title'],
                            'meta': f'{r["category"] or "חומר הסברה"} · {r["created_at"][:10]}'})
    conn.close()
    return jsonify(results)


# ---------------------------------------------------------------- ניהול (admin)
@app.route('/api/users')
@roles_required('admin')
def list_users():
    conn = get_db()
    rows = [dict(r) for r in conn.execute(
        'SELECT id, name, username, role, active, created_at FROM users ORDER BY id')]
    conn.close()
    return jsonify(rows)


@app.route('/api/users', methods=['POST'])
@roles_required('admin')
def create_user():
    d = request.get_json(force=True)
    if not d.get('name') or not d.get('username') or not d.get('password'):
        return jsonify({'error': 'שם, שם משתמש וסיסמה חובה'}), 400
    if d.get('role') not in ROLES:
        return jsonify({'error': 'תפקיד לא חוקי'}), 400
    if len(d['password']) < 8:
        return jsonify({'error': 'סיסמה חייבת לפחות 8 תווים'}), 400
    conn = get_db()
    try:
        cur = conn.execute('INSERT INTO users (name, username, password_hash, role, active, created_at) '
                           'VALUES (?,?,?,?,1,?)',
                           (d['name'].strip(), d['username'].strip(),
                            generate_password_hash(d['password']), d['role'], now()))
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({'error': 'שם המשתמש כבר קיים'}), 400
    log_action(conn, 'create_user', 'user', cur.lastrowid)
    conn.commit()
    conn.close()
    return jsonify({'id': cur.lastrowid})


@app.route('/api/users/<int:uid>', methods=['PUT'])
@roles_required('admin')
def update_user(uid):
    d = request.get_json(force=True)
    conn = get_db()
    target = conn.execute('SELECT * FROM users WHERE id=?', (uid,)).fetchone()
    if not target:
        conn.close()
        return jsonify({'error': 'משתמש לא נמצא'}), 404
    # הגנה על מנהל המערכת האחרון
    if target['role'] == 'admin' and (d.get('role', 'admin') != 'admin' or d.get('active') == 0):
        admins = conn.execute("SELECT COUNT(*) c FROM users WHERE role='admin' AND active=1").fetchone()['c']
        if admins <= 1:
            conn.close()
            return jsonify({'error': 'לא ניתן להסיר את מנהל המערכת האחרון'}), 400
    sets, params = [], []
    for f in ('name', 'role', 'active'):
        if f in d:
            if f == 'role' and d[f] not in ROLES:
                conn.close()
                return jsonify({'error': 'תפקיד לא חוקי'}), 400
            sets.append(f'{f}=?')
            params.append(d[f])
    if d.get('password'):
        if len(d['password']) < 8:
            conn.close()
            return jsonify({'error': 'סיסמה חייבת לפחות 8 תווים'}), 400
        sets.append('password_hash=?')
        params.append(generate_password_hash(d['password']))
    if sets:
        params.append(uid)
        conn.execute(f'UPDATE users SET {", ".join(sets)} WHERE id=?', params)
        log_action(conn, 'update_user', 'user', uid)
        conn.commit()
    conn.close()
    return jsonify({'ok': True})


def _list_table(table):
    conn = get_db()
    rows = [dict(r) for r in conn.execute(
        f'SELECT * FROM {table} ORDER BY sort_order, name')]
    conn.close()
    return jsonify(rows)


def _add_to_table(table):
    name = (request.get_json(force=True).get('name') or '').strip()
    if not name:
        return jsonify({'error': 'שם חובה'}), 400
    conn = get_db()
    try:
        conn.execute(f'INSERT INTO {table} (name, active, sort_order) VALUES (?,1,999)', (name,))
        conn.commit()
    except sqlite3.IntegrityError:
        conn.close()
        return jsonify({'error': 'כבר קיים'}), 400
    conn.close()
    return jsonify({'ok': True})


def _delete_from_table(table, item_id):
    conn = get_db()
    conn.execute(f'DELETE FROM {table} WHERE id=?', (item_id,))
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/topics')
def topics():
    return _list_table('topics')


@app.route('/api/topics', methods=['POST'])
@roles_required('admin', 'lead')
def add_topic():
    return _add_to_table('topics')


@app.route('/api/topics/<int:item_id>', methods=['DELETE'])
@roles_required('admin', 'lead')
def delete_topic(item_id):
    return _delete_from_table('topics', item_id)


@app.route('/api/authorities')
def authorities():
    return _list_table('authorities')


@app.route('/api/authorities', methods=['POST'])
@roles_required('admin', 'lead')
def add_authority():
    return _add_to_table('authorities')


@app.route('/api/authorities/<int:item_id>', methods=['DELETE'])
@roles_required('admin', 'lead')
def delete_authority(item_id):
    return _delete_from_table('authorities', item_id)


@app.route('/api/settings/status')
@roles_required('admin')
def settings_status():
    key = get_api_key()
    conn = get_db()
    key_from_db = bool(get_setting(conn, 'anthropic_api_key'))
    conn.close()
    from urllib.parse import quote
    tok_enc = quote(WEBHOOK_TOKEN, safe='') if WEBHOOK_TOKEN else ''
    return jsonify({
        'ai_enabled': ai_enabled(),
        'key_masked': (key[:12] + '…' + key[-4:]) if key and len(key) > 20 else None,
        'key_source': 'settings' if key_from_db else ('env' if ANTHROPIC_API_KEY else None),
        'model': get_model(),
        'webhook_configured': bool(WEBHOOK_TOKEN),
        'webhook_url': (request.url_root.rstrip('/') + '/api/webhook/whatsapp?token=' + tok_enc)
        if WEBHOOK_TOKEN else None,
        'greenapi_url': (request.url_root.rstrip('/') + '/api/webhook/greenapi?token=' + tok_enc)
        if WEBHOOK_TOKEN else None,
        'wa_send_enabled': wa_send_enabled(),
        'greenapi_instance': _greenapi_creds()[0],
        'is_prod': IS_PROD,
    })


@app.route('/api/settings/greenapi', methods=['POST'])
@roles_required('admin')
def set_greenapi():
    """הגדרת חיבור Green API לשליחת וואטסאפ. ריק = ניתוק."""
    d = request.get_json(force=True)
    iid = (d.get('instance_id') or '').strip()
    tok = (d.get('token') or '').strip()
    conn = get_db()
    if iid and tok:
        set_setting(conn, 'greenapi_instance_id', iid)
        set_setting(conn, 'greenapi_token', tok)
        log_action(conn, 'set_greenapi', 'settings', None, 'חיבור וואטסאפ עודכן')
    else:
        conn.execute("DELETE FROM settings WHERE key IN ('greenapi_instance_id','greenapi_token')")
        log_action(conn, 'clear_greenapi', 'settings', None)
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'wa_send_enabled': wa_send_enabled()})


@app.route('/api/settings/api-key', methods=['POST'])
@roles_required('admin')
def set_api_key():
    """שמירת מפתח Anthropic בהגדרות (מסך המנהל). ריק = מחיקה וחזרה למשתנה סביבה."""
    key = (request.get_json(force=True).get('key') or '').strip()
    if key and not key.startswith('sk-ant-'):
        return jsonify({'error': 'מפתח לא תקין — אמור להתחיל ב-sk-ant'}), 400
    conn = get_db()
    if key:
        set_setting(conn, 'anthropic_api_key', key)
        log_action(conn, 'set_api_key', 'settings', None, 'מפתח API עודכן')
    else:
        conn.execute("DELETE FROM settings WHERE key='anthropic_api_key'")
        log_action(conn, 'clear_api_key', 'settings', None)
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'ai_enabled': ai_enabled()})


@app.route('/api/settings/model', methods=['POST'])
@roles_required('admin')
def set_model():
    m = (request.get_json(force=True).get('model') or '').strip()
    conn = get_db()
    set_setting(conn, 'anthropic_model', m or None)
    conn.commit()
    conn.close()
    return jsonify({'ok': True})


@app.route('/api/settings/wa-log')
@roles_required('admin')
def wa_webhook_log():
    """יומן אירועי וואטסאפ אחרונים — לאבחון קליטה."""
    conn = get_db()
    rows = [dict(r) for r in conn.execute(
        "SELECT detail, created_at FROM audit_log WHERE action='wa_webhook' "
        'ORDER BY id DESC LIMIT 20')]
    conn.close()
    return jsonify(rows)


@app.route('/api/backup')
@roles_required('admin')
def backup():
    """גיבוי מלא: DB + כל הקבצים שהועלו, כקובץ ZIP להורדה."""
    # העתקה בטוחה של ה-DB (גם תוך כדי כתיבה) דרך sqlite backup API
    tmp_db = os.path.join(DATA_ROOT, '.backup_db_tmp.sqlite')
    src = sqlite3.connect(DB_PATH)
    dst = sqlite3.connect(tmp_db)
    with dst:
        src.backup(dst)
    dst.close()
    src.close()
    zip_path = os.path.join(DATA_ROOT, '.backup_latest.zip')
    with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as z:
        z.write(tmp_db, 'hasbara.db')
        for name in os.listdir(UPLOAD_DIR):
            fp = os.path.join(UPLOAD_DIR, name)
            if os.path.isfile(fp):
                z.write(fp, 'uploads/' + name)
    os.remove(tmp_db)
    conn = get_db()
    log_action(conn, 'backup_download', 'system', None,
               f'{os.path.getsize(zip_path) // 1024}KB')
    conn.commit()
    conn.close()
    return send_file(zip_path, download_name=f'gibui-hasbara-{today()}.zip',
                     as_attachment=True)


@app.route('/api/audit')
@roles_required('admin')
def audit():
    conn = get_db()
    rows = []
    for r in conn.execute('SELECT * FROM audit_log ORDER BY id DESC LIMIT 200'):
        d = dict(r)
        d['user_name'] = user_name(conn, r['user_id'])
        rows.append(d)
    conn.close()
    return jsonify(rows)


# ---------------------------------------------------------------- static
@app.route('/')
def index():
    return send_from_directory(BASE_DIR, 'index.html')


# אין catch-all סטטי בכוונה — רק index.html ו-/static מוגשים (לא חושפים קוד מקור)

init_db()

if __name__ == '__main__':
    app.run(debug=not IS_PROD, port=5080)
